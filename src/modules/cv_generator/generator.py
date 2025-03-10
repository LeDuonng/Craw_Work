#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Module tạo CV tự động
"""

import os
import json
import openai
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from src.utils.config import OPENAI_API_KEY, OUTPUT_DIR
from src.utils.logger import get_logger
from src.utils.helper import geocode_address

# Khởi tạo logger
logger = get_logger()

# Khởi tạo OpenAI API
openai.api_key = OPENAI_API_KEY

class CVGenerator:
    """
    Class tạo CV tự động
    """
    
    def __init__(self):
        """
        Khởi tạo generator
        """
        # Kiểm tra API key
        if not OPENAI_API_KEY:
            logger.warning("OpenAI API key không được cấu hình")
        
        # Tạo thư mục output nếu chưa tồn tại
        self.output_dir = os.path.join(OUTPUT_DIR, "cv")
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_cv_content(self, user_info):
        """
        Tạo nội dung CV
        
        Args:
            user_info (dict): Thông tin người dùng
        
        Returns:
            dict: Nội dung CV
        """
        try:
            # Kiểm tra thông tin người dùng
            if not user_info or not isinstance(user_info, dict):
                logger.error("Thông tin người dùng không hợp lệ")
                return None
            
            # Kiểm tra các trường bắt buộc
            required_fields = ['full_name', 'email', 'phone']
            for field in required_fields:
                if field not in user_info or not user_info[field]:
                    logger.error(f"Thiếu thông tin bắt buộc: {field}")
                    return None
            
            # Chuẩn bị prompt cho OpenAI
            prompt = self._prepare_prompt(user_info)
            
            # Gọi OpenAI API
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Bạn là một chuyên gia viết CV chuyên nghiệp. Hãy tạo nội dung CV dựa trên thông tin được cung cấp."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )
            
            # Xử lý kết quả
            cv_content = response.choices[0].message.content
            
            # Parse kết quả thành các phần
            sections = self._parse_cv_content(cv_content)
            
            # Thêm thông tin người dùng
            sections['user_info'] = user_info
            
            logger.info(f"Đã tạo nội dung CV cho {user_info['full_name']}")
            
            return sections
        
        except Exception as e:
            logger.error(f"Lỗi khi tạo nội dung CV: {e}")
            return None
    
    def _prepare_prompt(self, user_info):
        """
        Chuẩn bị prompt cho OpenAI
        
        Args:
            user_info (dict): Thông tin người dùng
        
        Returns:
            str: Prompt
        """
        prompt = f"""
        Hãy tạo nội dung CV chuyên nghiệp dựa trên thông tin sau:
        
        THÔNG TIN CÁ NHÂN:
        - Họ và tên: {user_info.get('full_name', '')}
        - Email: {user_info.get('email', '')}
        - Số điện thoại: {user_info.get('phone', '')}
        - Địa chỉ: {user_info.get('address', '')}
        - Ngày sinh: {user_info.get('birth_date', '')}
        
        KINH NGHIỆM LÀM VIỆC:
        {user_info.get('experience', 'Không có thông tin')}
        
        HỌC VẤN:
        {user_info.get('education', 'Không có thông tin')}
        
        KỸ NĂNG:
        {user_info.get('skills', 'Không có thông tin')}
        
        THÔNG TIN BỔ SUNG:
        {user_info.get('additional_info', 'Không có thông tin')}
        
        Hãy tạo nội dung CV với các phần sau:
        1. Tóm tắt bản thân (Summary)
        2. Kinh nghiệm làm việc (Experience)
        3. Học vấn (Education)
        4. Kỹ năng (Skills)
        5. Chứng chỉ (Certifications) - nếu có
        6. Dự án (Projects) - nếu có
        7. Thông tin bổ sung (Additional Information) - nếu có
        
        Hãy viết nội dung chuyên nghiệp, súc tích và phù hợp với ngành nghề của ứng viên. Mỗi phần hãy bắt đầu bằng tiêu đề "## [Tên phần]" để dễ dàng phân tách.
        """
        
        return prompt
    
    def _parse_cv_content(self, content):
        """
        Parse nội dung CV thành các phần
        
        Args:
            content (str): Nội dung CV
        
        Returns:
            dict: Các phần của CV
        """
        sections = {}
        
        # Tách nội dung thành các phần
        parts = content.split('##')
        
        for part in parts:
            part = part.strip()
            if not part:
                continue
            
            # Tách tiêu đề và nội dung
            lines = part.split('\n', 1)
            if len(lines) < 2:
                continue
            
            title = lines[0].strip()
            content = lines[1].strip()
            
            # Lưu vào dictionary
            key = title.lower().replace(' ', '_')
            sections[key] = content
        
        return sections
    
    def export_to_docx(self, cv_content, filename=None):
        """
        Xuất CV ra file DOCX
        
        Args:
            cv_content (dict): Nội dung CV
            filename (str): Tên file
        
        Returns:
            str: Đường dẫn đến file
        """
        try:
            # Kiểm tra nội dung CV
            if not cv_content or not isinstance(cv_content, dict):
                logger.error("Nội dung CV không hợp lệ")
                return None
            
            # Tạo tên file nếu chưa có
            if not filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"CV_{cv_content['user_info']['full_name'].replace(' ', '_')}_{timestamp}.docx"
            
            # Đường dẫn đầy đủ
            filepath = os.path.join(self.output_dir, filename)
            
            # Tạo document
            doc = Document()
            
            # Thiết lập style
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            # Thêm thông tin cá nhân
            user_info = cv_content['user_info']
            
            # Tiêu đề
            title = doc.add_heading(user_info['full_name'], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thông tin liên hệ
            contact_info = doc.add_paragraph()
            contact_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_info.add_run(f"Email: {user_info['email']} | Điện thoại: {user_info['phone']}")
            if 'address' in user_info and user_info['address']:
                contact_info.add_run(f" | Địa chỉ: {user_info['address']}")
            
            # Thêm các phần nội dung
            for key, content in cv_content.items():
                if key == 'user_info':
                    continue
                
                # Tiêu đề phần
                title = key.replace('_', ' ').title()
                doc.add_heading(title, level=1)
                
                # Nội dung phần
                doc.add_paragraph(content)
            
            # Lưu file
            doc.save(filepath)
            
            logger.info(f"Đã xuất CV ra file {filepath}")
            
            return filepath
        
        except Exception as e:
            logger.error(f"Lỗi khi xuất CV ra file DOCX: {e}")
            return None
    
    def export_to_pdf(self, cv_content, filename=None, photo_path=None):
        """
        Xuất CV ra file PDF
        
        Args:
            cv_content (dict): Nội dung CV
            filename (str): Tên file
            photo_path (str): Đường dẫn đến ảnh
        
        Returns:
            str: Đường dẫn đến file
        """
        try:
            # Kiểm tra nội dung CV
            if not cv_content or not isinstance(cv_content, dict):
                logger.error("Nội dung CV không hợp lệ")
                return None
            
            # Tạo tên file nếu chưa có
            if not filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"CV_{cv_content['user_info']['full_name'].replace(' ', '_')}_{timestamp}.pdf"
            
            # Đường dẫn đầy đủ
            filepath = os.path.join(self.output_dir, filename)
            
            # Tạo document
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            
            # Tạo style
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='Center', alignment=1))
            
            # Tạo danh sách các phần tử
            elements = []
            
            # Thêm thông tin cá nhân
            user_info = cv_content['user_info']
            
            # Tiêu đề
            elements.append(Paragraph(user_info['full_name'], styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Thông tin liên hệ
            contact_text = f"Email: {user_info['email']} | Điện thoại: {user_info['phone']}"
            if 'address' in user_info and user_info['address']:
                contact_text += f" | Địa chỉ: {user_info['address']}"
            elements.append(Paragraph(contact_text, styles['Center']))
            elements.append(Spacer(1, 12))
            
            # Thêm ảnh nếu có
            if photo_path and os.path.exists(photo_path):
                img = Image(photo_path, width=100, height=100)
                elements.append(img)
                elements.append(Spacer(1, 12))
            
            # Thêm các phần nội dung
            for key, content in cv_content.items():
                if key == 'user_info':
                    continue
                
                # Tiêu đề phần
                title = key.replace('_', ' ').title()
                elements.append(Paragraph(title, styles['Heading1']))
                elements.append(Spacer(1, 6))
                
                # Nội dung phần
                paragraphs = content.split('\n')
                for p in paragraphs:
                    if p.strip():
                        elements.append(Paragraph(p, styles['Normal']))
                        elements.append(Spacer(1, 6))
                
                elements.append(Spacer(1, 12))
            
            # Tạo PDF
            doc.build(elements)
            
            logger.info(f"Đã xuất CV ra file {filepath}")
            
            return filepath
        
        except Exception as e:
            logger.error(f"Lỗi khi xuất CV ra file PDF: {e}")
            return None 