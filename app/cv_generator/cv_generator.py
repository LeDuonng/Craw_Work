"""
CV Generator module
"""
import os
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from app.utils.openai_helper import generate_cv_with_openai
from app.utils.config import CV_OUTPUT_DIR

# Thiết lập logger
logger = logging.getLogger(__name__)

class CVGenerator:
    """
    Lớp tạo CV từ thông tin người dùng và thông tin công việc
    """
    
    def __init__(self):
        """
        Khởi tạo CVGenerator
        """
        # Đảm bảo thư mục đầu ra tồn tại
        os.makedirs(CV_OUTPUT_DIR, exist_ok=True)
        logger.info(f"Khởi tạo CV Generator, thư mục đầu ra: {CV_OUTPUT_DIR}")
    
    def generate_cv_content(self, user_info, job_info):
        """
        Tạo nội dung CV sử dụng OpenAI
        
        Args:
            user_info (dict): Thông tin cá nhân của người dùng
            job_info (dict): Thông tin về công việc đang ứng tuyển
            
        Returns:
            str: Nội dung CV được tạo
        """
        logger.info("Bắt đầu tạo nội dung CV với OpenAI")
        content = generate_cv_with_openai(user_info, job_info)
        logger.info("Đã tạo xong nội dung CV")
        return content
    
    def generate_docx(self, cv_content, filename):
        """
        Tạo file CV định dạng DOCX
        
        Args:
            cv_content (str): Nội dung CV
            filename (str): Tên file đầu ra (không bao gồm phần mở rộng)
            
        Returns:
            str: Đường dẫn đến file DOCX đã tạo
        """
        try:
            logger.info(f"Bắt đầu tạo file DOCX: {filename}.docx")
            
            # Tạo document mới
            doc = Document()
            
            # Thiết lập style
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            # Thiết lập margin
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Phân tích nội dung CV
            lines = cv_content.split('\n')
            
            # Thêm nội dung vào document
            for line in lines:
                line = line.strip()
                if not line:
                    # Thêm dòng trống
                    doc.add_paragraph()
                elif line.startswith('# '):
                    # Tiêu đề lớn
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line[2:])
                    run.bold = True
                    run.font.size = Pt(16)
                elif line.startswith('## '):
                    # Tiêu đề vừa
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.bold = True
                    run.font.size = Pt(14)
                elif line.startswith('### '):
                    # Tiêu đề nhỏ
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    run.bold = True
                    run.font.size = Pt(12)
                elif line.startswith('- '):
                    # Danh sách không thứ tự
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('* '):
                    # Danh sách không thứ tự (dạng khác)
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Danh sách có thứ tự
                    p = doc.add_paragraph(line[3:], style='List Number')
                else:
                    # Văn bản thông thường
                    p = doc.add_paragraph(line)
            
            # Lưu file
            output_path = os.path.join(CV_OUTPUT_DIR, f"{filename}.docx")
            doc.save(output_path)
            
            logger.info(f"Đã tạo thành công file DOCX: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Lỗi khi tạo file DOCX: {e}")
            print(f"Lỗi khi tạo file DOCX: {e}")
            return None
    
    def generate_pdf(self, cv_content, filename):
        """
        Tạo file CV định dạng PDF
        
        Args:
            cv_content (str): Nội dung CV
            filename (str): Tên file đầu ra (không bao gồm phần mở rộng)
            
        Returns:
            str: Đường dẫn đến file PDF đã tạo
        """
        try:
            logger.info(f"Bắt đầu tạo file PDF: {filename}.pdf")
            
            # Đường dẫn đầu ra
            output_path = os.path.join(CV_OUTPUT_DIR, f"{filename}.pdf")
            
            # Tạo document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Thiết lập styles
            styles = getSampleStyleSheet()
            
            # Sửa lỗi: Kiểm tra xem style 'Title' đã tồn tại trong styles chưa
            # Nếu đã tồn tại, sử dụng style đó thay vì tạo mới
            if 'Title' not in styles:
                styles.add(ParagraphStyle(
                    name='Title',
                    parent=styles['Heading1'],
                    fontSize=16,
                    alignment=1  # Center
                ))
            
            # Tạo các style khác nếu chưa tồn tại
            custom_styles = {
                'Heading2': {
                    'parent': 'Heading2',
                    'fontSize': 14
                },
                'Heading3': {
                    'parent': 'Heading3',
                    'fontSize': 12
                },
                'Normal': {
                    'parent': 'Normal',
                    'fontSize': 11,
                    'leading': 14
                },
                'Bullet': {
                    'parent': 'Normal',
                    'fontSize': 11,
                    'leftIndent': 20,
                    'firstLineIndent': -15,
                    'leading': 14
                }
            }
            
            # Thêm các style tùy chỉnh nếu chưa tồn tại
            for style_name, properties in custom_styles.items():
                if style_name not in styles:
                    styles.add(ParagraphStyle(
                        name=style_name,
                        parent=styles[properties['parent']],
                        fontSize=properties['fontSize'],
                        **{k:v for k,v in properties.items() if k not in ['parent', 'fontSize']}
                    ))
            
            # Phân tích nội dung CV
            lines = cv_content.split('\n')
            
            # Tạo danh sách các phần tử để thêm vào document
            elements = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Thêm khoảng trống
                    elements.append(Spacer(1, 0.1 * inch))
                elif line.startswith('# '):
                    # Tiêu đề lớn
                    elements.append(Paragraph(line[2:], styles['Title']))
                    elements.append(Spacer(1, 0.1 * inch))
                elif line.startswith('## '):
                    # Tiêu đề vừa
                    elements.append(Paragraph(line[3:], styles['Heading2']))
                    elements.append(Spacer(1, 0.1 * inch))
                elif line.startswith('### '):
                    # Tiêu đề nhỏ
                    elements.append(Paragraph(line[4:], styles['Heading3']))
                    elements.append(Spacer(1, 0.05 * inch))
                elif line.startswith('- '):
                    # Danh sách không thứ tự
                    elements.append(Paragraph('• ' + line[2:], styles['Bullet']))
                elif line.startswith('* '):
                    # Danh sách không thứ tự (dạng khác)
                    elements.append(Paragraph('• ' + line[2:], styles['Bullet']))
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Danh sách có thứ tự
                    num = line.split('.')[0]
                    text = line[len(num) + 2:]
                    elements.append(Paragraph(f"{num}. {text}", styles['Bullet']))
                else:
                    # Văn bản thông thường
                    elements.append(Paragraph(line, styles['Normal']))
            
            # Xây dựng document
            doc.build(elements)
            
            logger.info(f"Đã tạo thành công file PDF: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Lỗi khi tạo file PDF: {e}")
            print(f"Lỗi khi tạo file PDF: {e}")
            return None 